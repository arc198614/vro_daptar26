from flask import Flask, render_template, request, redirect, url_for, flash, jsonify
import os
from datetime import datetime
from gsheet_helper import GSheetHelper

app = Flask(__name__)
app.secret_key = os.urandom(24)

# Configuration
SHEET_ID = os.environ.get('GOOGLE_SHEET_ID', '11NCa_DbttL6x4Fq_oHRFooweNfNPNPq6nIlPl7OUQVU')
gsheet = None

def get_gsheet():
    global gsheet
    if gsheet is None:
        try:
            gsheet = GSheetHelper(SHEET_ID)
        except Exception as e:
            print(f"Error connecting to Google Sheets: {e}")
    return gsheet

@app.route('/')
def index():
    """Dashboard showing summary of inspections."""
    helper = get_gsheet()
    inspections = []
    if helper:
        inspections = helper.get_sheet_data('Inspections!A:E')
    
    return render_template('index.html', title="डॅशबोर्ड - ग्रा.म.अ. दप्तर तपासणी", inspections=inspections)

@app.route('/inspect', methods=['GET', 'POST'])
def inspect():
    """Form for new inspection."""
    helper = get_gsheet()

    # Load questions from 'Master_Q' sheet (needed for both GET and POST)
    questions = []
    if helper:
        questions = helper.get_sheet_data('Master_Q!A:D')

    # Fallback if sheet is empty or connection fails
    if not questions:
        questions = [
            {"ID": "1", "विभाग": "सामान्य", "प्रश्न": "दप्तर अद्ययावत आहे का?", "अपलोड आवश्यक": "हो"},
            {"ID": "2", "विभाग": "सामान्य", "प्रश्न": "नोंदवही पूर्ण आहे का?", "अपलोड आवश्यक": "नाही"}
        ]

    if request.method == 'POST':
        # Logic to save to Google Sheets
        data = request.form.to_dict()
        timestamp = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
        inspection_id = os.urandom(4).hex()
        
        # Handle file uploads and collect links
        import tempfile
        file_links = {}
        
        # Use system temp directory (works on Vercel/Lambda)
        upload_dir = tempfile.gettempdir()
            
        for key, file in request.files.items():
            if file and file.filename:
                file_path = os.path.join(upload_dir, file.filename)
                file.save(file_path)
                
                # Upload to Google Drive (folder_id is optional)
                # DRIVE_FOLDER_ID = 'your_folder_id_here' 
                _, link = helper.upload_file(file_path)
                if link:
                    file_links[key] = link
                
                # Clean up local file
                try:
                    os.remove(file_path)
                except:
                    pass

        # Prepare summary row for 'Inspections' sheet
        # Format: ID, Saja, Name, Registration Date, Date, Grade, File Link
        # Taking the first link as the primary file link for the summary
        primary_link = list(file_links.values())[0] if file_links else ""

        row = [
            inspection_id,
            data.get('saja_name'),
            data.get('vro_name'),
            data.get('registration_date'),  # Registration date
            timestamp,
            "Pending",
            primary_link
        ]

        if helper:
            helper.append_row('Inspections!A:G', row)
            
            # Save answers and remarks to 'Inspection_Answers' sheet for each question
            # First, get all question IDs from the questions list
            question_ids = [q['ID'] for q in questions] if questions else ['1', '2']

            for q_id in question_ids:
                answer = data.get(f'q_{q_id}', '')
                remark = data.get(f'remark_{q_id}', '').strip()

                # Save to Inspection_Answers sheet
                # Format: Inspection_ID, Question_ID, Answer, Remark
                answers_row = [
                    inspection_id,
                    q_id,
                    answer,
                    remark
                ]
                helper.append_row('Inspection_Answers!A:D', answers_row)

                # Also save to Compliance sheet if remark exists (for backward compatibility)
                if remark:
                    compliance_row = [
                        inspection_id,
                        remark, # अधिकारी शेरा (Remark)
                        "",     # वरिष्ठ मत
                        "",     # स्पष्टीकरण (GMA Explanation)
                        "Pending" # स्थिती (Status)
                    ]
                    helper.append_row('Compliance!A:E', compliance_row)

            print(f"Inspection {inspection_id} submitted with answers and remarks. File links: {file_links}")
            flash('तपासणी यशस्वीरित्या जतन केली आहे!', 'success')
        else:
            flash('त्रुटी: Google Sheet कनेक्शन अयशस्वी.', 'danger')
            
        return redirect(url_for('index'))

    return render_template('inspection_form.html', questions=questions, title="नवीन तपासणी")

@app.route('/compliance', methods=['GET', 'POST'])
def compliance():
    """Compliance tracking page."""
    helper = get_gsheet()
    
    if request.method == 'POST':
        # Update compliance logic
        log_id = request.form.get('log_id')
        remark = request.form.get('remark')
        explanation = request.form.get('explanation')
        senior_remark = request.form.get('senior_remark')
        status = request.form.get('status')
        
        # Updates list for Cols C, D, E (वरिष्ठ मत, स्पष्टीकरण, स्थिती)
        updates = [senior_remark, explanation, status]
        
        if helper and helper.update_compliance_row(log_id, remark, updates):
            flash('अनुपालन यशस्वीरित्या अपडेट केले आहे!', 'success')
        else:
            flash('त्रुटी: अपडेट अयशस्वी.', 'danger')
        return redirect(url_for('compliance'))

    compliance_data = []
    if helper:
        compliance_data = helper.get_sheet_data('Compliance!A:E')
    
    return render_template('compliance.html', title="अनुपालन अहवाल", compliance_data=compliance_data)

@app.route('/reports')
def reports():
    """Reports and statistics page."""
    helper = get_gsheet()
    inspections = []
    compliance = []
    if helper:
        inspections = helper.get_sheet_data('Inspections!A:G')
        compliance = helper.get_sheet_data('Compliance!A:E')
    
    # Simple stats for reports
    stats = {
        'total': len(inspections),
        'completed': len([i for i in inspections if i.get('एकूण ग्रेड') != 'Pending']),
        'pending_compliance': len([c for c in compliance if c.get('स्थिती') == 'Pending'])
    }
    
    return render_template('reports.html', title="अहवाल आणि सांख्यिकी", stats=stats, inspections=inspections)

@app.route('/export_pdf/<inspection_id>')
def export_pdf(inspection_id):
    """Generates a PDF report for a specific inspection."""
    helper = get_gsheet()
    if not helper:
        flash("Sheet connection error", "danger")
        return redirect(url_for('index'))

    # Fetch data
    inspections = helper.get_sheet_data('Inspections!A:G')
    inspection = next((i for i in inspections if i['ID'] == inspection_id), None)
    
    if not inspection:
        flash("Inspection not found", "danger")
        return redirect(url_for('index'))

    # Fetch related compliance/remarks
    compliance = helper.get_sheet_data('Compliance!A:E')
    remarks = [c for c in compliance if c['Log_ID'] == inspection_id]

    # Fetch inspection answers
    answers = helper.get_sheet_data('Inspection_Answers!A:D')
    inspection_answers = [a for a in answers if a['Inspection_ID'] == inspection_id]

    # Fetch questions from Master_Q to get actual question text
    questions = helper.get_sheet_data('Master_Q!A:D')
    question_map = {q['ID']: q['प्रश्न'] for q in questions}

    # PDF Generation with Unicode support using fpdf2
    try:
        from fpdf import FPDF

        # Create PDF with Unicode support
        pdf = FPDF()
        pdf.add_page()

        # Add Devanagari font if available
        base_dir = os.path.abspath(os.path.dirname(__file__))
        font_path = os.path.join(base_dir, 'NotoSansDevanagari-Regular.ttf')
        if os.path.isfile(font_path):
            pdf.add_font('NotoSansDevanagari', '', font_path, uni=True)
            pdf.set_font('NotoSansDevanagari', size=12)
        else:
            pdf.set_font('Arial', size=12)

        # Title - same as modal
        pdf.set_font(size=15)
        pdf.cell(0, 10, 'तपासणी अहवाल - GMA प्रणाली', 0, 1, 'C')
        pdf.ln(5)

        # Basic Information Table - same as modal
        pdf.set_font(size=12, style='B')
        pdf.cell(0, 10, txt="मूलभूत माहिती", ln=True)
        pdf.set_font(size=10)

        # Create basic info table
        pdf.cell(60, 8, txt="तपासणी ID:", border=1)
        pdf.cell(0, 8, txt=inspection.get('ID', ''), border=1, ln=True)

        pdf.cell(60, 8, txt="सजा/गाव:", border=1)
        pdf.cell(0, 8, txt=inspection.get('सजा', ''), border=1, ln=True)

        pdf.cell(60, 8, txt="ग्राम महसूल अधिकारी:", border=1)
        pdf.cell(0, 8, txt=inspection.get('नाव', ''), border=1, ln=True)

        pdf.cell(60, 8, txt="रुजू होण्याचा दिनांक:", border=1)
        pdf.cell(0, 8, txt=inspection.get('रुजू होण्याचा दिनांक', ''), border=1, ln=True)

        pdf.cell(60, 8, txt="तपासणी दिनांक:", border=1)
        pdf.cell(0, 8, txt=inspection.get('तारीख', ''), border=1, ln=True)

        pdf.cell(60, 8, txt="स्थिती:", border=1)
        status = inspection.get('एकूण ग्रेड', '')
        status_text = 'प्रलंबित' if status == 'Pending' else status
        pdf.cell(0, 8, txt=status_text, border=1, ln=True)

        pdf.cell(60, 8, txt="फाईल लिंक:", border=1)
        pdf.cell(0, 8, txt=inspection.get('फाईल लिंक', ''), border=1, ln=True)

        pdf.ln(10)

        # Inspection Answers Section - same as modal
        if inspection_answers:
            pdf.set_font(size=12, style='B')
            pdf.cell(0, 10, txt="तपासणी उत्तर", ln=True)
            pdf.set_font(size=10)

            # Table headers
            pdf.cell(15, 8, txt="प्रश्न क्र.", border=1, align='C')
            pdf.cell(50, 8, txt="प्रश्न", border=1, align='C')
            pdf.cell(30, 8, txt="उत्तर", border=1, align='C')
            pdf.cell(0, 8, txt="शेरा", border=1, align='C', ln=True)

            # Table data - use actual question text
            for answer in inspection_answers:
                question_id = str(answer.get('Question_ID', ''))
                question_text = question_map.get(question_id, f"प्रश्न {question_id}")
                pdf.cell(15, 8, txt=question_id, border=1)
                pdf.cell(50, 8, txt=question_text, border=1)
                pdf.cell(30, 8, txt=answer.get('Answer', ''), border=1)
                pdf.cell(0, 8, txt=answer.get('Remark', ''), border=1, ln=True)

            pdf.ln(5)

        # Documents Section
        if inspection.get('फाईल लिंक'):
            pdf.set_font(size=12, style='B')
            pdf.cell(0, 10, txt="अपलोड केलेली दस्तऐवज", ln=True)
            pdf.set_font(size=10)
            pdf.cell(0, 8, txt=f"दस्तऐवज लिंक: {inspection.get('फाईल लिंक', '')}", ln=True)
            pdf.ln(5)

        # Remarks & Compliance Section - same as modal
        if remarks:
            pdf.set_font(size=12, style='B')
            pdf.cell(0, 10, txt="अनुपालन आणि शेरा", ln=True)
            pdf.set_font(size=10)

            for remark in remarks:
                pdf.cell(0, 8, txt=f"शेरा: {remark.get('अधिकारी शेरा', '')}", ln=True)
                pdf.cell(0, 8, txt=f"वरिष्ठ मत: {remark.get('वरिष्ठ मत', '')}", ln=True)
                pdf.cell(0, 8, txt=f"स्पष्टीकरण: {remark.get('स्पष्टीकरण', '')}", ln=True)
                pdf.cell(0, 8, txt=f"स्थिती: {remark.get('स्थिती', '')}", ln=True)
                pdf.ln(3)
    except Exception as e:
        # Log the error and return a user‑friendly message
        print(f"PDF generation error: {e}")
        import traceback
        print(f"Traceback: {traceback.format_exc()}")
        flash('Failed to generate PDF report.', 'danger')
        return redirect(url_for('index'))

    from flask import Response
    try:
        pdf_bytes = pdf.output()
        response = Response(pdf_bytes, mimetype='application/pdf')
        response.headers.set('Content-Disposition', 'attachment', filename=f'Report_{inspection_id}.pdf')
        return response
    except Exception as e:
        print(f"PDF output error: {e}")
        flash('Error generating PDF file.', 'danger')
        return redirect(url_for('index'))

@app.route('/export_word/<inspection_id>')
def export_word(inspection_id):
    """Generates a Word document report for a specific inspection."""
    helper = get_gsheet()
    if not helper:
        flash("Sheet connection error", "danger")
        return redirect(url_for('index'))

    # Fetch data
    inspections = helper.get_sheet_data('Inspections!A:G')
    inspection = next((i for i in inspections if i['ID'] == inspection_id), None)

    if not inspection:
        flash("Inspection not found", "danger")
        return redirect(url_for('index'))

    # Fetch related compliance/remarks
    compliance = helper.get_sheet_data('Compliance!A:E')
    remarks = [c for c in compliance if c['Log_ID'] == inspection_id]

    # Fetch inspection answers
    answers = helper.get_sheet_data('Inspection_Answers!A:D')
    inspection_answers = [a for a in answers if a['Inspection_ID'] == inspection_id]

    # Fetch questions from Master_Q to get actual question text
    questions = helper.get_sheet_data('Master_Q!A:D')
    question_map = {q['ID']: q['प्रश्न'] for q in questions}

    # Word Document Generation with proper Unicode support
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from io import BytesIO
        import os

        doc = Document()

        # Add title with proper font
        title = doc.add_heading('तपासणी अहवाल - GMA प्रणाली', 0)
        title.style.font.size = Pt(16)

        # Basic Information Section
        doc.add_heading('मूलभूत माहिती', level=1)

        # Create table for basic info
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'

        # Header row
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = 'क्षेत्र'
        hdr_cells[1].text = 'माहिती'

        # Add data rows - same as modal display
        info_rows = [
            ('तपासणी ID', inspection.get('ID', '')),
            ('सजा/गाव', inspection.get('सजा', '')),
            ('ग्राम महसूल अधिकारी', inspection.get('नाव', '')),
            ('रुजू होण्याचा दिनांक', inspection.get('रुजू होण्याचा दिनांक', '')),
            ('तपासणी दिनांक', inspection.get('तारीख', '')),
            ('स्थिती', inspection.get('एकूण ग्रेड', '')),
            ('फाईल लिंक', inspection.get('फाईल लिंक', ''))
        ]

        for field_name, field_value in info_rows:
            row_cells = table.add_row().cells
            row_cells[0].text = field_name
            row_cells[1].text = field_value

        # Inspection Answers Section
        if inspection_answers:
            doc.add_heading('तपासणी उत्तर', level=1)

            answers_table = doc.add_table(rows=1, cols=4)
            answers_table.style = 'Table Grid'

            # Headers
            hdr_cells = answers_table.rows[0].cells
            hdr_cells[0].text = 'प्रश्न क्र.'
            hdr_cells[1].text = 'प्रश्न'
            hdr_cells[2].text = 'उत्तर'
            hdr_cells[3].text = 'शेरा'

            # Add answer rows - use actual question text
            for answer in inspection_answers:
                question_id = str(answer.get('Question_ID', ''))
                question_text = question_map.get(question_id, f"प्रश्न {question_id}")
                row_cells = answers_table.add_row().cells
                row_cells[0].text = question_id
                row_cells[1].text = question_text
                row_cells[2].text = answer.get('Answer', '')
                row_cells[3].text = answer.get('Remark', '')

        # Documents Section
        if inspection.get('फाईल लिंक'):
            doc.add_heading('अपलोड केलेली दस्तऐवज', level=1)
            doc.add_paragraph(f"दस्तऐवज लिंक: {inspection.get('फाईल लिंक', '')}")

        # Remarks & Compliance Section
        if remarks:
            doc.add_heading('अनुपालन आणि शेरा', level=1)

            for remark in remarks:
                doc.add_paragraph(f"शेरा: {remark.get('अधिकारी शेरा', '')}")
                doc.add_paragraph(f"वरिष्ठ मत: {remark.get('वरिष्ठ मत', '')}")
                doc.add_paragraph(f"स्पष्टीकरण: {remark.get('स्पष्टीकरण', '')}")
                doc.add_paragraph(f"स्थिती: {remark.get('स्थिती', '')}")
                doc.add_paragraph("")  # Empty line

        # Save to BytesIO
        doc_io = BytesIO()
        doc.save(doc_io)
        doc_io.seek(0)

    except Exception as e:
        print(f"Word document generation error: {e}")
        import traceback
        print(f"Traceback: {traceback.format_exc()}")
        flash('Failed to generate Word document report.', 'danger')
        return redirect(url_for('index'))

    from flask import Response
    response = Response(doc_io.getvalue(), mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response.headers.set('Content-Disposition', 'attachment', filename=f'Report_{inspection_id}.docx')
    return response

@app.route('/api/inspection/<inspection_id>')
def get_inspection_details(inspection_id):
    """API endpoint to get detailed inspection information."""
    helper = get_gsheet()
    if not helper:
        return {'error': 'Database connection failed'}, 500

    try:
        # Get inspection details
        inspections = helper.get_sheet_data('Inspections!A:G')
        inspection = next((i for i in inspections if i['ID'] == inspection_id), None)

        if not inspection:
            return {'error': 'Inspection not found'}, 404

        # Get inspection answers
        answers = helper.get_sheet_data('Inspection_Answers!A:D')
        inspection_answers = [a for a in answers if a['Inspection_ID'] == inspection_id]

        # Get compliance/remarks
        compliance = helper.get_sheet_data('Compliance!A:E')
        remarks = [c for c in compliance if c['Log_ID'] == inspection_id]

        # Get uploaded files (if any)
        # Note: Files are stored as links in the Inspections sheet under 'फाईल लिंक'
        files = []
        if inspection.get('फाईल लिंक'):
            files.append({
                'name': f"Inspection_{inspection_id}_file",
                'url': inspection['फाईल लिंक']
            })

        return jsonify({
            'inspection': inspection,
            'answers': inspection_answers,
            'remarks': remarks,
            'files': files
        })

    except Exception as e:
        print(f"Error fetching inspection details: {e}")
        return {'error': 'Internal server error'}, 500

if __name__ == '__main__':
    app.run(debug=True)
